import pandas as pd
from docx import Document

def get_current(a):
    if(a):
        return "Да"
    else:
        return"Нет"


# Параметры файлов
input_file = './input/output_table.xlsx'  # Файл Excel с отфильтрованными данными
output_file = './out/CWE_output.docx'  # Имя выходного Word файла

# Загрузка данных из Excel
df = pd.read_excel(input_file)


# Создание нового документа Word
doc = Document()

# Для каждой строки создаем таблицу с данными
for index, row in df.iterrows():
    # Добавляем заголовок таблицы
    doc.add_paragraph(f'Таблица {index + 233} - Паспорт уязвимости {str(row.get('Unnamed: 18', 'Не указано'))}')

    # Создаем таблицу с 2 столбцами
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'  # Устанавливаем стиль таблицы

    # Заполнение данных в таблицу
    data = [
        ("Идентификатор уязвимости",
         str(row.get('Unnamed: 18', 'Не указано')) + "\n" + str(row.get('Unnamed: 2', 'Не указано'))),
        ("Идентификатор типа ошибки", str(row.get('Unnamed: 24', 'Не указано'))),
        ("Базовый вектор уязвимости", str(row.get('Unnamed: 10', 'Не указано'))),
        ("Уровень опасности уязвимости", str(row.get('Unnamed: 12', 'Не указано'))),
        ("Статус уязвимости", str(row.get('Unnamed: 14', 'Не указано'))),
        ("Наличие эксплойта", str(row.get('Unnamed: 15', 'Не указано'))),
        ("Способ устранения", str(row.get('Unnamed: 22', 'Не указано'))),
        ("Участие в инцидентах", get_current(row.get('Unnamed: 20', 'Не указано'))),
    ]

    for label, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value) if pd.notna(value) else 'Не указано'

    # Добавляем разрыв строки после каждой таблицы
    doc.add_paragraph()

# Сохранение документа
doc.save(output_file)
print(f"Документ успешно сохранен в '{output_file}'")
