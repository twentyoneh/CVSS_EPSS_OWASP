import pandas as pd
import requests
from cvss import CVSS2, CVSS3, CVSS4
import OWAPS
from docx import Document
import math  # Для расчёта корня

def calculate_rms(values):
    """
    Вычисляет среднеквадратичное значение для нормализованных оценок.
    :param values: Список нормализованных значений.
    :return: Среднеквадратичное значение.
    """
    n = len(values)
    return math.sqrt(sum(v**2 for v in values) / n) if n > 0 else 0

def fetch_epss_data(cve_id):
    """
    Запрашивает данные EPSS по CVE.
    """
    api_url = f"https://api.first.org/data/v1/epss?cve={cve_id}"
    try:
        response = requests.get(api_url)
        response.raise_for_status()
        data = response.json()
        # Проверяем наличие данных в ключе 'data'
        if data and 'data' in data and isinstance(data['data'], list) and len(data['data']) > 0:
            epss_entry = data['data'][0]
            return {
                'epss': epss_entry.get('epss'),
                'percentile': epss_entry.get('percentile'),
                'date': epss_entry.get('date')
            }
        else:
            return {}
    except requests.exceptions.RequestException as e:
        print(f"Ошибка запроса API для {cve_id}: {e}")
        return {}


def normalize_score(s, smin, smax):
    """
    Нормализует значение оценки.
    :param s: Оценка (может быть строкой или None)
    :param smin: Минимальная оценка
    :param smax: Максимальная оценка
    :return: Нормализованная оценка
    """
    try:
        s = float(s)  # Преобразуем значение в float
        return (s - smin) / (smax - smin) if smax != smin else 0
    except (ValueError, TypeError):
        return 0  # Возвращаем 0 для некорректных или отсутствующих значений


def process_cve_data(input_excel, output_excel):
    """
    Обрабатывает данные из входного Excel файла, вызывает API для каждого CVE
    и сохраняет результаты в новый Excel файл и Word-документ.
    """
    metric_values = {'N': 0, 'L': 0.2, 'H': 0.5}
    # Читаем входной Excel файл
    df = pd.read_excel(input_excel)

    # Проверяем наличие столбца с CVE
    if 'CVE_filtered' not in df.columns:
        raise ValueError("Ожидается столбец 'CVE_filtered' в файле!")

    # Убираем строки без CVE
    df = df.dropna(subset=['CVE_filtered'])

    # Создаем список для хранения результатов
    results = []
    doc = Document()

    doc.add_heading('Результаты анализа CVSS3.1, CVSS4.0, EPSS, OWASP', level=1)

    # Создаем таблицу с 5 столбцами
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ['Уязвимость', 'Оценка CVSS 3.1', 'Оценка CVSS 4.0', 'Оценка EPSS', 'Оценка OWASP']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Переменные для нормализации
    epss_min, epss_max = 0, 1
    cvss_min, cvss_max = 0, 10

    # Запрос данных для каждого CVE
    for index, row in df.iterrows():
        cve_id = row['CVE_filtered']

        cvss4_score = CVSS4(row['CVSS4.0'])
        cvss3_score = CVSS3(f"CVSS:3.0/{row['Unnamed: 11']}")
        metrics = cvss3_score.metrics
        U_c = metric_values[metrics.get('C', 'N')]
        U_i = metric_values[metrics.get('I', 'N')]
        U_a = metric_values[metrics.get('A', 'N')]

        impact_score = 1 - (1 - U_c) * (1 - U_i) * (1 - U_a)

        epss_data = fetch_epss_data(cve_id)
        owaps_score = OWAPS.calculate_risk_owaps()

        result_row = {
            'CWE': row.get('Unnamed: 24', None),
            'CVE_filtered': cve_id,
            'CVSS3.1': row.get('Unnamed: 11', None),
            'CVSS4.0': row.get('CVSS4.0', None),
            'EPSS_Score': epss_data.get('epss', None),
            'EPSS_Percentile': epss_data.get('percentile', None),
            'Calc CVSS3.1': cvss3_score.base_score,
            'Calc CVSS4.0': cvss4_score.base_score,
            'Calc OWAPS': owaps_score,
            'Impact_Score': impact_score,
        }
        results.append(result_row)

        data = {
            'Уязвимость': cve_id,
            'Оценка CVSS 3.1': cvss3_score.base_score,
            'Оценка CVSS 4.0': cvss4_score.base_score,
            'Оценка EPSS': epss_data.get('epss', None),
            'Оценка OWASP': owaps_score,
        }
        row = table.add_row().cells
        row[0].text = str(data['Уязвимость'])
        row[1].text = str(data['Оценка CVSS 3.1'])
        row[2].text = str(data['Оценка CVSS 4.0'])
        row[3].text = str(data['Оценка EPSS'])
        row[4].text = str(data['Оценка OWASP'])

    # Добавляем таблицу нормализованных значений
    doc.add_paragraph()
    doc.add_heading('Нормализованные оценки критичности уязвимостей', level=1)

    norm_table = doc.add_table(rows=1, cols=5)
    norm_table.style = 'Table Grid'

    norm_headers = ['Уязвимость', 'Норм. CVSS 3.1', 'Норм. CVSS 4.0', 'Норм. EPSS', 'Норм. OWASP']
    for i, header in enumerate(norm_headers):
        norm_table.cell(0, i).text = header

    norm_values_list = []
    for res in results:
        # Нормализованные значения
        norm_values = [
            normalize_score(res['Calc CVSS3.1'], cvss_min, cvss_max),
            normalize_score(res['Calc CVSS4.0'], cvss_min, cvss_max),
            normalize_score(res.get('EPSS_Score', 0), epss_min, epss_max),
            normalize_score(res['Calc OWAPS'], cvss_min, cvss_max)
        ]

        row = norm_table.add_row().cells
        row[0].text = str(res['CVE_filtered'])
        row[1].text = f"{norm_values[0]:.2f}"
        row[2].text = f"{norm_values[1]:.2f}"
        row[3].text = f"{norm_values[2]:.5f}"
        row[4].text = f"{norm_values[3]:.2f}"

        norm_values_list.append(norm_values)

    # Добавляем таблицу среднеквадратичных значений
    doc.add_paragraph()
    doc.add_heading('Среднеквадратичная оценка критичности уязвимостей', level=1)

    rms_table = doc.add_table(rows=1, cols=2)
    rms_table.style = 'Table Grid'

    rms_headers = ['Уязвимость', 'Среднеквадратичная оценка']
    for i, header in enumerate(rms_headers):
        rms_table.cell(0, i).text = header

    rms_values = []
    for res in results:
        # Нормализованные значения
        norm_values = [
            normalize_score(res['Calc CVSS3.1'], cvss_min, cvss_max),
            normalize_score(res['Calc CVSS4.0'], cvss_min, cvss_max),
            normalize_score(res.get('EPSS_Score', 0), epss_min, epss_max),
            normalize_score(res['Calc OWAPS'], cvss_min, cvss_max)
        ]

        # Среднеквадратичное значение
        rms_score = calculate_rms(norm_values)

        # Добавление строки в таблицу
        row = rms_table.add_row().cells
        row[0].text = str(res['CVE_filtered'])
        row[1].text = f"{rms_score:.2f}"

        rms_values.append(rms_score)

    # Теперь считаем вероятность успеха для каждой уязвимости
    doc.add_paragraph()
    doc.add_heading('Вероятность успеха реализации единичной атаки', level=1)

    success_table = doc.add_table(rows=1, cols=2)
    success_table.style = 'Table Grid'

    success_headers = ['Уязвимость', 'Вероятность успеха']
    for i, header in enumerate(success_headers):
        success_table.cell(0, i).text = header

    for rms_score in rms_values:
        # Суммируем все среднеквадратичные значения
        total_rms_score = sum(rms_values)

        # Рассчитываем вероятность успеха для каждой уязвимости
        success_probability = rms_score / total_rms_score

        # Добавляем строку в таблицу
        row = success_table.add_row().cells
        row[0].text = str(results[rms_values.index(rms_score)]['CVE_filtered'])
        row[1].text = f"{success_probability:.5f}"

    output_file = './out/vulnerability_analysis.docx'
    doc.save(output_file)

    print(f"Документ успешно сохранен как {output_file}")
    # Преобразуем список результатов в DataFrame
    results_df = pd.DataFrame(results)

    # Сохраняем в новый Excel файл
    results_df.to_excel(output_excel, index=False)
    print(f"Результаты успешно сохранены в '{output_excel}'")
